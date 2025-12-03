"""
Servicios para procesamiento de documentos y búsqueda semántica (RAG)
"""
import os
import re
import math
import logging
from typing import List, Dict, Tuple, Optional
from django.utils import timezone
from django.db.models import Q

logger = logging.getLogger(__name__)


# ============ EXTRACTOR DE TEXTO ============

class ExtractorTexto:
    """Extrae texto de diferentes formatos de archivo"""

    @staticmethod
    def extraer(archivo_path: str, tipo_mime: str = '') -> str:
        """
        Extrae texto de un archivo según su tipo.
        Soporta: PDF, DOCX, TXT
        """
        extension = os.path.splitext(archivo_path)[1].lower()

        try:
            if extension == '.pdf' or 'pdf' in tipo_mime:
                return ExtractorTexto._extraer_pdf(archivo_path)
            elif extension in ['.docx', '.doc'] or 'word' in tipo_mime:
                return ExtractorTexto._extraer_docx(archivo_path)
            elif extension == '.txt' or 'text/plain' in tipo_mime:
                return ExtractorTexto._extraer_txt(archivo_path)
            else:
                logger.warning(f"Tipo de archivo no soportado: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extrayendo texto de {archivo_path}: {e}")
            raise

    @staticmethod
    def _extraer_pdf(archivo_path: str) -> str:
        """Extrae texto de un PDF usando PyPDF2"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(archivo_path)
            texto = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    texto.append(page_text)
            return '\n\n'.join(texto)
        except ImportError:
            logger.error("PyPDF2 no está instalado. Instalar con: pip install PyPDF2")
            raise ImportError("Se requiere PyPDF2 para procesar PDFs")

    @staticmethod
    def _extraer_docx(archivo_path: str) -> str:
        """Extrae texto de un DOCX usando python-docx"""
        try:
            from docx import Document
            doc = Document(archivo_path)
            texto = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texto.append(para.text)
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        texto.append(row_text)
            return '\n\n'.join(texto)
        except ImportError:
            logger.error("python-docx no está instalado. Instalar con: pip install python-docx")
            raise ImportError("Se requiere python-docx para procesar archivos Word")

    @staticmethod
    def _extraer_txt(archivo_path: str) -> str:
        """Extrae texto de un archivo TXT"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(archivo_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"No se pudo decodificar el archivo {archivo_path}")


# ============ FRAGMENTADOR DE TEXTO ============

class FragmentadorTexto:
    """Divide texto en fragmentos para embeddings"""

    def __init__(
        self,
        tamaño_fragmento: int = 500,
        solapamiento: int = 50,
        separadores: List[str] = None
    ):
        """
        Args:
            tamaño_fragmento: Número aproximado de palabras por fragmento
            solapamiento: Palabras de solapamiento entre fragmentos
            separadores: Lista de separadores para dividir (por defecto: párrafos, oraciones)
        """
        self.tamaño_fragmento = tamaño_fragmento
        self.solapamiento = solapamiento
        self.separadores = separadores or ['\n\n', '\n', '. ', '? ', '! ']

    def fragmentar(self, texto: str) -> List[Dict]:
        """
        Divide el texto en fragmentos con metadatos.

        Returns:
            Lista de diccionarios con: contenido, num_caracteres, num_tokens_aprox
        """
        if not texto or not texto.strip():
            return []

        # Limpiar texto
        texto = self._limpiar_texto(texto)

        # Dividir primero por párrafos
        parrafos = self._dividir_por_separador(texto, '\n\n')

        fragmentos = []
        fragmento_actual = []
        palabras_actuales = 0

        for parrafo in parrafos:
            palabras_parrafo = len(parrafo.split())

            # Si el párrafo es muy grande, dividirlo más
            if palabras_parrafo > self.tamaño_fragmento:
                # Guardar lo que tenemos
                if fragmento_actual:
                    fragmentos.append(' '.join(fragmento_actual))
                    fragmento_actual = []
                    palabras_actuales = 0

                # Dividir párrafo largo
                sub_fragmentos = self._dividir_parrafo_largo(parrafo)
                fragmentos.extend(sub_fragmentos)
            else:
                # Verificar si cabe en el fragmento actual
                if palabras_actuales + palabras_parrafo <= self.tamaño_fragmento:
                    fragmento_actual.append(parrafo)
                    palabras_actuales += palabras_parrafo
                else:
                    # Guardar fragmento actual y empezar nuevo
                    if fragmento_actual:
                        fragmentos.append(' '.join(fragmento_actual))

                    # Solapamiento: tomar últimas palabras del fragmento anterior
                    if self.solapamiento > 0 and fragmento_actual:
                        ultimo = fragmento_actual[-1]
                        palabras_ultimo = ultimo.split()
                        solapamiento = ' '.join(palabras_ultimo[-self.solapamiento:])
                        fragmento_actual = [solapamiento, parrafo]
                        palabras_actuales = self.solapamiento + palabras_parrafo
                    else:
                        fragmento_actual = [parrafo]
                        palabras_actuales = palabras_parrafo

        # Guardar último fragmento
        if fragmento_actual:
            fragmentos.append(' '.join(fragmento_actual))

        # Crear diccionarios con metadatos
        resultado = []
        for i, contenido in enumerate(fragmentos):
            if contenido.strip():
                resultado.append({
                    'numero': i,
                    'contenido': contenido.strip(),
                    'num_caracteres': len(contenido),
                    'num_tokens_aprox': len(contenido.split())
                })

        return resultado

    def _limpiar_texto(self, texto: str) -> str:
        """Limpia el texto de caracteres innecesarios"""
        # Normalizar espacios en blanco
        texto = re.sub(r'\s+', ' ', texto)
        # Restaurar saltos de párrafo
        texto = re.sub(r' ?\n ?', '\n', texto)
        texto = re.sub(r'\n{3,}', '\n\n', texto)
        # Eliminar caracteres de control
        texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', texto)
        return texto.strip()

    def _dividir_por_separador(self, texto: str, separador: str) -> List[str]:
        """Divide texto por un separador, filtrando vacíos"""
        partes = texto.split(separador)
        return [p.strip() for p in partes if p.strip()]

    def _dividir_parrafo_largo(self, parrafo: str) -> List[str]:
        """Divide un párrafo largo en fragmentos más pequeños"""
        fragmentos = []
        oraciones = re.split(r'(?<=[.!?])\s+', parrafo)

        fragmento_actual = []
        palabras_actuales = 0

        for oracion in oraciones:
            palabras_oracion = len(oracion.split())

            if palabras_actuales + palabras_oracion <= self.tamaño_fragmento:
                fragmento_actual.append(oracion)
                palabras_actuales += palabras_oracion
            else:
                if fragmento_actual:
                    fragmentos.append(' '.join(fragmento_actual))
                fragmento_actual = [oracion]
                palabras_actuales = palabras_oracion

        if fragmento_actual:
            fragmentos.append(' '.join(fragmento_actual))

        return fragmentos


# ============ GENERADOR DE EMBEDDINGS ============

class GeneradorEmbeddings:
    """
    Genera embeddings usando sentence-transformers.
    Usa el modelo multilingüe para español.
    """

    _modelo = None
    MODELO_DEFAULT = 'paraphrase-multilingual-MiniLM-L12-v2'

    @classmethod
    def obtener_modelo(cls):
        """Carga el modelo (singleton para eficiencia)"""
        if cls._modelo is None:
            try:
                from sentence_transformers import SentenceTransformer
                logger.info(f"Cargando modelo de embeddings: {cls.MODELO_DEFAULT}")
                cls._modelo = SentenceTransformer(cls.MODELO_DEFAULT)
                logger.info("Modelo cargado exitosamente")
            except ImportError:
                logger.error("sentence-transformers no está instalado")
                raise ImportError(
                    "Se requiere sentence-transformers. "
                    "Instalar con: pip install sentence-transformers"
                )
        return cls._modelo

    @classmethod
    def generar_embedding(cls, texto: str) -> List[float]:
        """Genera embedding para un texto"""
        modelo = cls.obtener_modelo()
        embedding = modelo.encode(texto, convert_to_numpy=True)
        return embedding.tolist()

    @classmethod
    def generar_embeddings_batch(cls, textos: List[str]) -> List[List[float]]:
        """Genera embeddings para múltiples textos (más eficiente)"""
        if not textos:
            return []
        modelo = cls.obtener_modelo()
        embeddings = modelo.encode(textos, convert_to_numpy=True, show_progress_bar=False)
        return [e.tolist() for e in embeddings]


# ============ BUSCADOR SEMÁNTICO ============

class BuscadorSemantico:
    """Realiza búsqueda semántica usando similitud coseno"""

    @staticmethod
    def similitud_coseno(vec1: List[float], vec2: List[float]) -> float:
        """Calcula similitud coseno entre dos vectores"""
        if not vec1 or not vec2 or len(vec1) != len(vec2):
            return 0.0

        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        norm1 = math.sqrt(sum(a * a for a in vec1))
        norm2 = math.sqrt(sum(b * b for b in vec2))

        if norm1 == 0 or norm2 == 0:
            return 0.0

        return dot_product / (norm1 * norm2)

    @classmethod
    def buscar(
        cls,
        query: str,
        usuario,
        empresa_id: str = None,
        top_k: int = 5,
        umbral_similitud: float = 0.3
    ) -> List[Dict]:
        """
        Busca fragmentos relevantes para una consulta.

        Args:
            query: Texto de búsqueda
            usuario: Usuario que realiza la búsqueda (para permisos)
            empresa_id: Filtrar por empresa específica
            top_k: Número máximo de resultados
            umbral_similitud: Similitud mínima para incluir resultado

        Returns:
            Lista de fragmentos con su similitud
        """
        from .models import Documento, FragmentoDocumento, NivelAcceso

        # Generar embedding de la query
        try:
            query_embedding = GeneradorEmbeddings.generar_embedding(query)
        except Exception as e:
            logger.error(f"Error generando embedding para query: {e}")
            return []

        # Construir filtro de documentos accesibles
        filtro_docs = Q(activo=True, procesado=True)

        # Filtro por empresa
        if empresa_id:
            filtro_docs &= Q(empresa_id=empresa_id) | Q(empresa__isnull=True)
        else:
            # Documentos globales o de empresas del usuario
            empresas_usuario = usuario.empresas.all()
            filtro_docs &= Q(empresa__isnull=True) | Q(empresa__in=empresas_usuario)

        # Filtro por nivel de acceso
        if usuario.rol == 'administrador':
            pass  # Ve todo
        elif usuario.rol == 'empleador':
            filtro_docs &= Q(tipo_acceso__in=[
                NivelAcceso.PUBLICO, NivelAcceso.EMPRESA, NivelAcceso.RRHH
            ]) | Q(created_by=usuario)
        else:
            filtro_docs &= Q(tipo_acceso__in=[
                NivelAcceso.PUBLICO, NivelAcceso.EMPRESA
            ]) | Q(created_by=usuario)

        # Obtener documentos accesibles
        docs_accesibles = Documento.objects.filter(filtro_docs).values_list('id', flat=True)

        # Obtener fragmentos con embeddings
        fragmentos = FragmentoDocumento.objects.filter(
            documento_id__in=docs_accesibles,
            embedding__isnull=False
        ).select_related('documento')

        # Calcular similitud para cada fragmento
        resultados = []
        for fragmento in fragmentos:
            similitud = cls.similitud_coseno(query_embedding, fragmento.embedding)
            if similitud >= umbral_similitud:
                resultados.append({
                    'fragmento_id': str(fragmento.id),
                    'documento_id': str(fragmento.documento_id),
                    'documento_titulo': fragmento.documento.titulo,
                    'documento_tipo': fragmento.documento.tipo,
                    'contenido': fragmento.contenido,
                    'similitud': round(similitud, 4),
                    'numero_fragmento': fragmento.numero_fragmento,
                })

        # Ordenar por similitud y limitar
        resultados.sort(key=lambda x: x['similitud'], reverse=True)
        return resultados[:top_k]


# ============ PROCESADOR DE DOCUMENTOS ============

class ProcesadorDocumentos:
    """Orquesta el procesamiento completo de documentos"""

    def __init__(
        self,
        tamaño_fragmento: int = 500,
        solapamiento: int = 50
    ):
        self.fragmentador = FragmentadorTexto(
            tamaño_fragmento=tamaño_fragmento,
            solapamiento=solapamiento
        )

    def procesar_documento(self, documento) -> Dict:
        """
        Procesa un documento: extrae texto, fragmenta y genera embeddings.

        Args:
            documento: Instancia de Documento

        Returns:
            Diccionario con resultado del procesamiento
        """
        from .models import FragmentoDocumento

        try:
            # 1. Extraer texto del archivo
            archivo_path = documento.archivo.path
            logger.info(f"Procesando documento: {documento.titulo}")

            texto = ExtractorTexto.extraer(archivo_path, documento.tipo_mime)
            if not texto:
                raise ValueError("No se pudo extraer texto del documento")

            # Guardar texto completo
            documento.contenido_texto = texto

            # 2. Fragmentar texto
            fragmentos_data = self.fragmentador.fragmentar(texto)
            logger.info(f"Documento fragmentado en {len(fragmentos_data)} partes")

            if not fragmentos_data:
                raise ValueError("El documento no produjo fragmentos válidos")

            # 3. Generar embeddings en batch
            textos = [f['contenido'] for f in fragmentos_data]
            embeddings = GeneradorEmbeddings.generar_embeddings_batch(textos)

            # 4. Eliminar fragmentos anteriores
            FragmentoDocumento.objects.filter(documento=documento).delete()

            # 5. Crear nuevos fragmentos
            fragmentos_creados = []
            for i, (frag_data, embedding) in enumerate(zip(fragmentos_data, embeddings)):
                fragmento = FragmentoDocumento.objects.create(
                    documento=documento,
                    contenido=frag_data['contenido'],
                    numero_fragmento=i,
                    embedding=embedding,
                    num_tokens=frag_data['num_tokens_aprox'],
                    num_caracteres=frag_data['num_caracteres']
                )
                fragmentos_creados.append(fragmento)

            # 6. Marcar como procesado
            documento.procesado = True
            documento.fecha_procesado = timezone.now()
            documento.error_procesamiento = ''
            documento.save()

            logger.info(f"Documento {documento.titulo} procesado exitosamente")

            return {
                'success': True,
                'documento_id': str(documento.id),
                'texto_extraido': len(texto),
                'fragmentos_creados': len(fragmentos_creados),
            }

        except Exception as e:
            logger.error(f"Error procesando documento {documento.id}: {e}")
            documento.procesado = False
            documento.error_procesamiento = str(e)
            documento.save()

            return {
                'success': False,
                'documento_id': str(documento.id),
                'error': str(e)
            }

    def reprocesar_todos(self, empresa_id: str = None) -> Dict:
        """
        Reprocesa todos los documentos (o los de una empresa).
        Útil al cambiar configuración de fragmentación.
        """
        from .models import Documento

        filtro = Q(activo=True)
        if empresa_id:
            filtro &= Q(empresa_id=empresa_id)

        documentos = Documento.objects.filter(filtro)
        resultados = {
            'total': documentos.count(),
            'exitosos': 0,
            'fallidos': 0,
            'errores': []
        }

        for doc in documentos:
            resultado = self.procesar_documento(doc)
            if resultado['success']:
                resultados['exitosos'] += 1
            else:
                resultados['fallidos'] += 1
                resultados['errores'].append({
                    'documento': doc.titulo,
                    'error': resultado.get('error', 'Error desconocido')
                })

        return resultados


# ============ FUNCIONES DE UTILIDAD ============

def buscar_en_documentos(
    query: str,
    usuario,
    empresa_id: str = None,
    top_k: int = 5
) -> List[Dict]:
    """
    Función de conveniencia para búsqueda semántica.
    """
    return BuscadorSemantico.buscar(
        query=query,
        usuario=usuario,
        empresa_id=empresa_id,
        top_k=top_k
    )


def obtener_contexto_para_ia(
    query: str,
    usuario,
    empresa_id: str = None,
    max_fragmentos: int = 3
) -> str:
    """
    Obtiene contexto relevante de documentos para usar en prompts de IA.
    Formatea los fragmentos encontrados como contexto.
    """
    resultados = buscar_en_documentos(
        query=query,
        usuario=usuario,
        empresa_id=empresa_id,
        top_k=max_fragmentos
    )

    if not resultados:
        return ""

    lineas = ["INFORMACIÓN RELEVANTE DE DOCUMENTOS DE LA EMPRESA:\n"]
    for i, r in enumerate(resultados, 1):
        lineas.append(f"[Fuente: {r['documento_titulo']}]")
        lineas.append(r['contenido'])
        lineas.append("")

    return "\n".join(lineas)
