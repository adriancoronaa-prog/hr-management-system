"""
Generador de documentos Word para el sistema RRHH
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import date
from typing import Dict


class GeneradorDocumentos:
    """Genera documentos Word para RRHH"""

    @classmethod
    def generar_constancia_laboral(cls, empleado, incluir_salario: bool = False) -> BytesIO:
        """Genera constancia laboral en Word"""
        doc = Document()

        # Configurar margenes
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        # Membrete empresa
        empresa = empleado.empresa
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(empresa.razon_social.upper())
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"RFC: {empresa.rfc}")

        direccion = cls._obtener_direccion_empresa(empresa)
        if direccion:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(direccion)

        doc.add_paragraph()

        # Titulo
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONSTANCIA LABORAL")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        # Contenido
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        fecha_ingreso_texto = empleado.fecha_ingreso.strftime('%d de %B de %Y') if empleado.fecha_ingreso else 'N/A'

        texto = f"A quien corresponda:\n\n"
        texto += f"Por medio de la presente, hacemos constar que {empleado.nombre_completo.upper()}, "
        texto += f"con CURP {empleado.curp or 'N/A'} y RFC {empleado.rfc or 'N/A'}, "
        texto += f"labora en esta empresa desde el {fecha_ingreso_texto} "
        texto += f"desempenando el puesto de {empleado.puesto or 'N/A'} en el departamento de {empleado.departamento or 'N/A'}."

        if incluir_salario and empleado.salario_mensual:
            texto += f"\n\nPercibe un salario mensual de ${empleado.salario_mensual:,.2f} (pesos mexicanos)."

        texto += f"\n\nSe extiende la presente constancia a peticion del interesado "
        texto += f"para los fines legales que a su derecho convengan, "
        texto += f"en la Ciudad de Mexico, a {date.today().strftime('%d de %B de %Y')}."

        p.add_run(texto)

        doc.add_paragraph()
        doc.add_paragraph()

        # Firma
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 40)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Recursos Humanos")
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(empresa.razon_social)

        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @classmethod
    def generar_carta_recomendacion(cls, empleado, destinatario: str = "A quien corresponda") -> BytesIO:
        """Genera carta de recomendacion"""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        empresa = empleado.empresa

        # Fecha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Ciudad de Mexico, {date.today().strftime('%d de %B de %Y')}")

        doc.add_paragraph()

        # Destinatario
        p = doc.add_paragraph()
        run = p.add_run(destinatario)
        run.bold = True
        p = doc.add_paragraph()
        p.add_run("Presente")

        doc.add_paragraph()

        # Contenido
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if empleado.fecha_ingreso:
            antiguedad = (date.today() - empleado.fecha_ingreso).days // 365
            anos_texto = f"{antiguedad} ano{'s' if antiguedad != 1 else ''}" if antiguedad > 0 else "menos de un ano"
        else:
            anos_texto = "un tiempo"

        texto = f"Por medio de la presente, me permito recomendar ampliamente a "
        texto += f"{empleado.nombre_completo.upper()}, quien labora en {empresa.razon_social} "
        texto += f"durante {anos_texto}, desempenandose como {empleado.puesto or 'colaborador'}.\n\n"

        texto += f"Durante su estancia en nuestra empresa, ha demostrado ser una persona responsable, "
        texto += f"comprometida y con excelente actitud de servicio. Ha cumplido satisfactoriamente "
        texto += f"con las funciones asignadas a su puesto, mostrando profesionalismo y dedicacion.\n\n"

        texto += f"Por lo anterior, no tengo inconveniente en recomendarlo para cualquier "
        texto += f"oportunidad laboral que se le presente, seguro de que sera un valioso elemento "
        texto += f"para cualquier organizacion."

        p.add_run(texto)

        doc.add_paragraph()
        doc.add_paragraph()

        # Despedida
        p = doc.add_paragraph()
        p.add_run("Sin mas por el momento, quedo a sus ordenes.")

        doc.add_paragraph()
        doc.add_paragraph()

        # Firma
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Atentamente,")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 40)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Recursos Humanos")
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(empresa.razon_social)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @classmethod
    def generar_carta_amonestacion(cls, empleado, motivo: str, fecha_incidente: date = None) -> BytesIO:
        """Genera carta de amonestacion"""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        empresa = empleado.empresa
        fecha_inc = fecha_incidente or date.today()

        # Titulo
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CARTA DE AMONESTACION")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Fecha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Ciudad de Mexico, {date.today().strftime('%d de %B de %Y')}")

        doc.add_paragraph()

        # Destinatario
        p = doc.add_paragraph()
        run = p.add_run(empleado.nombre_completo.upper())
        run.bold = True
        p = doc.add_paragraph()
        p.add_run(f"Puesto: {empleado.puesto or 'N/A'}")
        p = doc.add_paragraph()
        p.add_run(f"Departamento: {empleado.departamento or 'N/A'}")
        p = doc.add_paragraph()
        p.add_run("Presente")

        doc.add_paragraph()

        # Contenido
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        texto = f"Por medio de la presente, se le hace una amonestacion formal "
        texto += f"debido al siguiente motivo:\n\n"
        texto += f"{motivo}\n\n"
        texto += f"Este incidente ocurrio el dia {fecha_inc.strftime('%d de %B de %Y')}.\n\n"
        texto += f"Le exhortamos a corregir esta conducta y apegarse al reglamento interno de trabajo. "
        texto += f"De reincidir en esta falta, se tomaran las medidas disciplinarias correspondientes "
        texto += f"conforme a la Ley Federal del Trabajo y al reglamento de la empresa."

        p.add_run(texto)

        doc.add_paragraph()
        doc.add_paragraph()

        # Firmas
        table = doc.add_table(rows=2, cols=2)

        cell1 = table.cell(0, 0)
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 25 + "\nRecursos Humanos")

        cell2 = table.cell(0, 1)
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 25 + "\nEmpleado (Acuse de recibo)")

        cell1 = table.cell(1, 0)
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n\n_" * 25 + "\nTestigo")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @classmethod
    def generar_contrato_laboral(cls, empleado, datos_contrato: Dict = None) -> BytesIO:
        """Genera contrato laboral basico"""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        empresa = empleado.empresa
        datos = datos_contrato or {}

        # Titulo
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONTRATO INDIVIDUAL DE TRABAJO")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Intro
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        direccion = cls._obtener_direccion_empresa(empresa)

        texto = f"Contrato individual de trabajo que celebran por una parte {empresa.razon_social}, "
        texto += f"con RFC {empresa.rfc}, representada en este acto por su representante legal, "
        texto += f"a quien en lo sucesivo se le denominara 'EL PATRON', y por la otra parte "
        texto += f"{empleado.nombre_completo.upper()}, con CURP {empleado.curp or 'N/A'} "
        texto += f"y RFC {empleado.rfc or 'N/A'}, a quien en lo sucesivo se le denominara 'EL TRABAJADOR', "
        texto += f"al tenor de las siguientes:"
        p.add_run(texto)

        doc.add_paragraph()

        # Clausulas
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CLAUSULAS")
        run.bold = True

        salario_diario = empleado.salario_diario or 0
        salario_texto = cls._numero_a_letras(float(salario_diario))
        fecha_ingreso_texto = empleado.fecha_ingreso.strftime('%d de %B de %Y') if empleado.fecha_ingreso else date.today().strftime('%d de %B de %Y')

        clausulas = [
            f"PRIMERA.- EL TRABAJADOR se obliga a prestar sus servicios personales subordinados "
            f"a EL PATRON, consistentes en el puesto de {empleado.puesto or 'N/A'}, en el departamento de {empleado.departamento or 'N/A'}.",

            f"SEGUNDA.- La duracion de este contrato es por tiempo {datos.get('tipo', 'indeterminado')}.",

            f"TERCERA.- EL TRABAJADOR prestara sus servicios en las instalaciones de EL PATRON "
            f"ubicadas en {direccion or 'las oficinas de la empresa'}.",

            f"CUARTA.- La jornada de trabajo sera de {datos.get('jornada', 'lunes a viernes de 9:00 a 18:00 horas')}, "
            f"con una hora para tomar alimentos.",

            f"QUINTA.- EL PATRON pagara a EL TRABAJADOR un salario diario de "
            f"${salario_diario:,.2f} ({salario_texto} pesos), "
            f"que sera cubierto de manera quincenal.",

            f"SEXTA.- EL TRABAJADOR disfrutara de un periodo de vacaciones conforme a lo establecido "
            f"en la Ley Federal del Trabajo, asi como del pago de aguinaldo y demas prestaciones de ley.",

            f"SEPTIMA.- EL TRABAJADOR se obliga a cumplir con el Reglamento Interior de Trabajo "
            f"de EL PATRON.",

            f"OCTAVA.- Para todo lo no previsto en este contrato, las partes se sujetaran a las "
            f"disposiciones de la Ley Federal del Trabajo.",
        ]

        for clausula in clausulas:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(clausula)

        doc.add_paragraph()

        # Cierre
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto = f"Leido que fue el presente contrato, y enteradas las partes de su contenido y alcance legal, "
        texto += f"lo firman por duplicado en la Ciudad de Mexico, a {fecha_ingreso_texto}."
        p.add_run(texto)

        doc.add_paragraph()
        doc.add_paragraph()

        # Firmas
        table = doc.add_table(rows=3, cols=2)

        cell1 = table.cell(0, 0)
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 25)

        cell2 = table.cell(0, 1)
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 25)

        cell1 = table.cell(1, 0)
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("EL PATRON")
        run.bold = True

        cell2 = table.cell(1, 1)
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("EL TRABAJADOR")
        run.bold = True

        cell1 = table.cell(2, 0)
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(empresa.razon_social)

        cell2 = table.cell(2, 1)
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(empleado.nombre_completo)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @classmethod
    def _obtener_direccion_empresa(cls, empresa) -> str:
        """Construye la direccion de la empresa"""
        partes = []
        if hasattr(empresa, 'calle') and empresa.calle:
            partes.append(empresa.calle)
        if hasattr(empresa, 'numero_exterior') and empresa.numero_exterior:
            partes.append(f"No. {empresa.numero_exterior}")
        if hasattr(empresa, 'colonia') and empresa.colonia:
            partes.append(f"Col. {empresa.colonia}")
        if hasattr(empresa, 'codigo_postal') and empresa.codigo_postal:
            partes.append(f"C.P. {empresa.codigo_postal}")
        if hasattr(empresa, 'municipio') and empresa.municipio:
            partes.append(empresa.municipio)
        if hasattr(empresa, 'estado') and empresa.estado:
            partes.append(empresa.estado)

        return ", ".join(partes) if partes else ""

    @staticmethod
    def _numero_a_letras(numero: float) -> str:
        """Convierte numero a texto (simplificado)"""
        partes = str(int(numero)).split('.')
        entero = int(partes[0])

        unidades = ['', 'un', 'dos', 'tres', 'cuatro', 'cinco', 'seis', 'siete', 'ocho', 'nueve']
        decenas = ['', 'diez', 'veinte', 'treinta', 'cuarenta', 'cincuenta', 'sesenta', 'setenta', 'ochenta', 'noventa']
        centenas = ['', 'ciento', 'doscientos', 'trescientos', 'cuatrocientos', 'quinientos', 'seiscientos', 'setecientos', 'ochocientos', 'novecientos']

        if entero == 0:
            return 'cero'
        if entero == 100:
            return 'cien'

        resultado = []

        if entero >= 1000:
            miles = entero // 1000
            if miles == 1:
                resultado.append('mil')
            else:
                resultado.append(f'{unidades[miles]} mil')
            entero %= 1000

        if entero >= 100:
            resultado.append(centenas[entero // 100])
            entero %= 100

        if entero >= 10:
            if entero < 20:
                especiales = {10: 'diez', 11: 'once', 12: 'doce', 13: 'trece', 14: 'catorce', 15: 'quince', 16: 'dieciseis', 17: 'diecisiete', 18: 'dieciocho', 19: 'diecinueve'}
                resultado.append(especiales.get(entero, ''))
                entero = 0
            else:
                resultado.append(decenas[entero // 10])
                entero %= 10
                if entero > 0:
                    resultado.append('y')

        if entero > 0:
            resultado.append(unidades[entero])

        return ' '.join(resultado)
